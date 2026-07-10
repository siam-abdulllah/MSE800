#!/usr/bin/env python3
"""Generate Feasibility Study Word document (GeeksforGeeks 8-type framework)."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Feasibility_Study.docx")
GEEKS_URL = (
    "https://www.geeksforgeeks.org/software-engineering/"
    "types-of-feasibility-study-in-software-project-development/"
)


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return table


def build_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    h = doc.add_heading("Feasibility Study", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local Community Skill Exchange and Help Board")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed  |  MSE800 Assessment 2")
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "This feasibility study follows the eight-type framework from GeeksforGeeks: "
        "Technical, Operational, Economic, Legal, Schedule, Cultural & Political, "
        "Market, and Resource feasibility."
    )
    doc.add_paragraph(f"Reference: {GEEKS_URL}")

    doc.add_heading("1. Need of Feasibility Study", level=2)
    doc.add_paragraph(
        "A feasibility study determines whether the project is practically feasible, "
        "identifies risks, narrows alternatives, and provides a go/no-go conclusion "
        "before development begins."
    )

    doc.add_heading("2. Aim of Feasibility Study", level=2)
    add_table(doc, ["Aim", "Assessment"],
        [
            ["System contributes to objectives", "Yes — connects neighbours, builds trust"],
            ["Implementable with current technology", "Yes — Django, PostgreSQL, Redis, Channels"],
            ["Integrates with existing systems", "Yes — REST API, SMTP email; future portal integration"],
        ],
    )

    doc.add_heading("3. Feasibility Study Process", level=2)
    add_table(doc, ["Step", "Activity"],
        [
            ["1", "Information assessment — define concept and objectives"],
            ["2", "Information collection — gather tech, cost, schedule, legal, market data"],
            ["3", "Report writing — document eight feasibility types"],
            ["4", "General information — summary and recommendation"],
        ],
    )

    doc.add_heading("4. Eight Types of Feasibility — Summary", level=2)
    add_table(doc, ["Type", "Verdict", "Priority"],
        [
            ["Technical", "Feasible", "High"],
            ["Operational", "Feasible", "High"],
            ["Economic", "Feasible", "Most important"],
            ["Legal", "Feasible", "Moderate"],
            ["Schedule", "Feasible", "High"],
            ["Cultural & Political", "Feasible", "Moderate"],
            ["Market", "Feasible", "Moderate"],
            ["Resource", "Feasible", "High"],
            ["Overall", "Proceed with development", ""],
        ],
    )

    doc.add_page_break()
    doc.add_heading("5. Technical Feasibility", level=2)
    doc.add_paragraph(
        "Analyses hardware, software, team skills, and technology maintainability. "
        "Django, PostgreSQL, Redis, and Channels are mature. Team has MSE800 experience. "
        "Suburb/postcode search fallback reduces GeoDjango risk. Verdict: Feasible."
    )

    doc.add_heading("6. Operational Feasibility", level=2)
    doc.add_paragraph(
        "Analyses service delivery, usability, and post-deployment maintenance. "
        "Clear roles (Member, Moderator, Admin); responsive UI; Django admin for maintenance. "
        "Verdict: Feasible."
    )

    doc.add_heading("7. Economic Feasibility", level=2)
    doc.add_paragraph(
        "Most important feasibility type (GeeksforGeeks). Analyses cost vs benefit."
    )
    add_table(doc, ["Cost / Benefit", "Amount / Value"],
        [
            ["Development cost", "$0 — assessment project"],
            ["Software licences", "$0 — open source"],
            ["Staging hosting", "$0–$15/month"],
            ["Community benefit", "Free local help exchange vs 10–20% commercial fees"],
            ["Learning outcome", "Full-stack software engineering demonstration"],
        ],
    )
    doc.add_paragraph("Verdict: Economically feasible — benefits outweigh costs.")

    doc.add_heading("8. Legal Feasibility", level=2)
    doc.add_paragraph(
        "Privacy Act compliance; suburb/postcode only (no exact addresses); hashed passwords; "
        "report/block moderation; open-source licence compliance; no payments (out of scope). "
        "Verdict: Legally feasible."
    )

    doc.add_heading("9. Schedule Feasibility", level=2)
    add_table(doc, ["Sprint", "Weeks", "Focus"],
        [
            ["Sprint 1", "1–2", "Foundation — auth, profiles, listings"],
            ["Sprint 2", "3–4", "Core features — search, chat, ratings"],
            ["Release Sprint", "5–6", "UAT, deployment, regression"],
        ],
    )
    doc.add_paragraph(
        "2 developers × ~20 hrs/week × 6 weeks ≈ 240 person-hours. "
        "Verdict: Schedule feasible."
    )

    doc.add_page_break()
    doc.add_heading("10. Cultural and Political Feasibility", level=2)
    doc.add_paragraph(
        "Platform supports neighbourhood mutual aid — aligns with community cooperation values. "
        "Moderation handles cultural sensitivity. No political barriers for assessment scope. "
        "Verdict: Feasible."
    )

    doc.add_heading("11. Market Feasibility", level=2)
    doc.add_paragraph(
        "Strong need for local help discovery. Differentiated from fee-based apps (Airtasker). "
        "Free community exchange niche is underserved. Verdict: Market feasible."
    )

    doc.add_heading("12. Resource Feasibility", level=2)
    add_table(doc, ["Resource", "Required", "Available"],
        [
            ["Developers", "2", "Group F — 2 members"],
            ["Budget", "$0–$20/month", "Free tiers / student credits"],
            ["Software", "Django, PostgreSQL, Redis", "Open source"],
            ["Time", "~240 person-hours", "6-week 3-sprint plan"],
        ],
    )
    doc.add_paragraph("Verdict: Resource feasible.")

    doc.add_heading("13. Feasibility Decision Matrix", level=2)
    add_table(doc, ["Type", "Weight", "Score (1–5)", "Weighted"],
        [
            ["Technical", "15%", "5", "0.75"],
            ["Operational", "12%", "4", "0.48"],
            ["Economic (most important)", "20%", "5", "1.00"],
            ["Legal", "8%", "4", "0.32"],
            ["Schedule", "15%", "4", "0.60"],
            ["Cultural & Political", "8%", "4", "0.32"],
            ["Market", "10%", "4", "0.40"],
            ["Resource", "12%", "5", "0.60"],
            ["Total", "100%", "", "4.47 / 5.00 — Proceed"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Recommendation: Proceed with development. All eight GeeksforGeeks feasibility types "
        "assessed as feasible. Prioritise Must requirements in Sprints 1–2; "
        "reserve Release Sprint for UAT and deployment."
    )

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
