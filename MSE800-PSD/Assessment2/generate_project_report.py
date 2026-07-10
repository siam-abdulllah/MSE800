#!/usr/bin/env python3
"""Generate full Project Report Word document including GeeksforGeeks feasibility study."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Report.docx")
GEEKS_URL = (
    "https://www.geeksforgeeks.org/software-engineering/"
    "types-of-feasibility-study-in-software-project-development/"
)


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return table


def build_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    h = doc.add_heading("Project Report", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local Community Skill Exchange and Help Board")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed  |  MSE800 Assessment 2")
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(
        "This report documents planning and feasibility analysis for a community help board "
        "platform. Feasibility follows the GeeksforGeeks eight-type framework. "
        "Overall score: 4.47/5.00. Recommendation: proceed with development."
    )

    doc.add_heading("2. Introduction", level=2)
    doc.add_paragraph(
        "Residents need local help but lack a trusted channel. This project delivers a web "
        "prototype for posting help listings, matching by location, chatting, and rating."
    )

    doc.add_heading("3. Technology Stack", level=2)
    add_table(doc, ["Layer", "Technology"],
        [
            ["Backend", "Python, Django, DRF"],
            ["Database", "PostgreSQL"],
            ["Real-time", "Django Channels + Redis"],
            ["Testing", "pytest-django"],
            ["Deployment", "Docker Compose"],
        ],
    )

    doc.add_heading("4. Agile Plan Summary", level=2)
    add_table(doc, ["Sprint", "Weeks", "Focus"],
        [
            ["Sprint 1", "1–2", "Auth, profiles, listings, admin"],
            ["Sprint 2", "3–4", "Search, chat, ratings, safety"],
            ["Release Sprint", "5–6", "UAT, deployment, regression"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("5. Feasibility Study", level=2)
    doc.add_paragraph(
        "Following GeeksforGeeks eight-type feasibility framework. "
        f"Reference: {GEEKS_URL}"
    )

    doc.add_heading("5.1 Need and Aim", level=3)
    doc.add_paragraph(
        "Need: go/no-go decision before development. "
        "Aim: confirm objectives alignment, current technology suitability, and integration potential."
    )

    doc.add_heading("5.2 Eight Feasibility Types — Summary", level=3)
    add_table(doc, ["Type", "Verdict", "Key Finding"],
        [
            ["Technical", "Feasible", "Django stack mature; team skills sufficient"],
            ["Operational", "Feasible", "Clear workflows; Django admin maintenance"],
            ["Economic", "Feasible", "Most important — ~$0 cost, high benefit"],
            ["Legal", "Feasible", "Privacy Act; moderation; open-source compliance"],
            ["Schedule", "Feasible", "6-week 3-sprint plan meets deadline"],
            ["Cultural & Political", "Feasible", "Mutual-aid values; moderation for sensitivity"],
            ["Market", "Feasible", "Local need; differentiated from fee-based apps"],
            ["Resource", "Feasible", "2 developers, tools, time available"],
        ],
        font_size=8,
    )

    doc.add_heading("5.3 Technical Feasibility", level=3)
    doc.add_paragraph(
        "Hardware, software, team skills, maintainability. Django + PostgreSQL + Redis + Channels. "
        "Verdict: Feasible."
    )

    doc.add_heading("5.4 Operational Feasibility", level=3)
    doc.add_paragraph(
        "Service delivery, usability, post-deployment maintenance. Member/Moderator/Admin roles. "
        "Verdict: Feasible."
    )

    doc.add_heading("5.5 Economic Feasibility (Most Important)", level=3)
    doc.add_paragraph(
        "Cost: ~$0–$20/month. Benefit: community connection, cost savings vs commercial apps, "
        "assessment learning. Verdict: Feasible."
    )

    doc.add_heading("5.6 Legal Feasibility", level=3)
    doc.add_paragraph(
        "Privacy Act, user content moderation, open-source licences. No payments in scope. "
        "Verdict: Feasible."
    )

    doc.add_heading("5.7 Schedule Feasibility", level=3)
    doc.add_paragraph(
        "6 weeks, 3 sprints, ~240 person-hours. Release Sprint for UAT. Verdict: Feasible."
    )

    doc.add_heading("5.8 Cultural, Market, and Resource Feasibility", level=3)
    doc.add_paragraph(
        "Cultural: aligns with community cooperation. Market: underserved free local exchange niche. "
        "Resource: human, financial, technology resources available. All feasible."
    )

    doc.add_heading("5.9 Decision Matrix", level=3)
    add_table(doc, ["Type", "Weight", "Score", "Weighted"],
        [
            ["Technical", "15%", "5", "0.75"],
            ["Operational", "12%", "4", "0.48"],
            ["Economic", "20%", "5", "1.00"],
            ["Legal", "8%", "4", "0.32"],
            ["Schedule", "15%", "4", "0.60"],
            ["Cultural & Political", "8%", "4", "0.32"],
            ["Market", "10%", "4", "0.40"],
            ["Resource", "12%", "5", "0.60"],
            ["Total", "100%", "", "4.47 / 5.00"],
        ],
    )

    doc.add_heading("6. Conclusion", level=2)
    doc.add_paragraph(
        "All eight GeeksforGeeks feasibility types assessed as feasible. "
        "Proceed with Sprint 1, Sprint 2, and Release Sprint."
    )

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
