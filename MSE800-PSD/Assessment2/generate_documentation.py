#!/usr/bin/env python3
"""Generate Assessment 2 project documentation (requirements + agile plan) as Word."""

import io
import math
import os
import sys

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "Local_Community_Skill_Exchange_Project_Documentation.docx",
)

DARK_BLUE = "#1B3A6B"
LIGHT_BLUE = "#5BA3E8"
YELLOW = "#E8B923"
RED = "#C0392B"
GEAR_COLOR = "#E8DDD0"
SPRINT_COLORS = {1: DARK_BLUE, 2: "#D4A017"}


FUNCTIONAL_REQUIREMENTS = [
    ("FR-01", "User registration", "New residents create accounts with email and password.", "Sprint 1", "accounts/views.py, accounts/models.py"),
    ("FR-02", "Email verification", "Verification link sent; unverified users cannot post or chat.", "Sprint 1", "accounts/views.py, notifications/email.py"),
    ("FR-03", "Login and logout", "Secure session-based authentication for members.", "Sprint 1", "accounts/views.py, accounts/urls.py"),
    ("FR-04", "Password reset", "Time-limited reset link via email.", "Sprint 1", "accounts/views.py, templates/emails/"),
    ("FR-05", "User profile", "Skills, bio, suburb/postcode, optional photo.", "Sprint 1", "profiles/models.py, profiles/views.py"),
    ("FR-06", "Role-based access", "Member, Moderator, and Admin roles with permissions.", "Sprint 1", "accounts/permissions.py, accounts/models.py"),
    ("FR-07", "Create help listing", "Post 'I need help' or 'I can help' with details.", "Sprint 1", "listings/models.py, listings/views.py"),
    ("FR-08", "Browse listings", "Public feed of active help listings.", "Sprint 1", "listings/views.py, templates/listings/"),
    ("FR-09", "Listing categories", "Predefined categories for discovery.", "Sprint 1", "listings/models.py (Category)"),
    ("FR-10", "Edit and close listing", "Owners edit or mark listings complete.", "Sprint 1", "listings/views.py"),
    ("FR-11", "Admin moderation panel", "Admins manage users, posts, and reports.", "Sprint 1", "moderation/admin.py, listings/admin.py"),
    ("FR-12", "Nearby search", "Filter listings by suburb, postcode, or radius.", "Sprint 2", "listings/filters.py, listings/services/location.py"),
    ("FR-13", "Skill-based matching", "Filter by skill keywords and category.", "Sprint 2", "listings/filters.py"),
    ("FR-14", "Express interest / match", "Open a private chat thread between users.", "Sprint 2", "listings/views.py, chat/models.py"),
    ("FR-15", "Real-time chat", "Live messaging via Django Channels + Redis.", "Sprint 2", "chat/consumers.py, config/asgi.py"),
    ("FR-16", "Chat history", "Persisted message history per conversation.", "Sprint 2", "chat/models.py, chat/views.py"),
    ("FR-17", "Star ratings", "1–5 star rating with optional review text.", "Sprint 2", "reviews/models.py, reviews/views.py"),
    ("FR-18", "Display aggregate rating", "Average rating shown on user profiles.", "Sprint 2", "profiles/views.py, reviews/models.py"),
    ("FR-19", "Report user or listing", "Reports queued for moderator review.", "Sprint 2", "moderation/models.py, moderation/views.py"),
    ("FR-20", "Block user", "Block prevents messages and hides listings.", "Sprint 2", "moderation/models.py, chat/services.py"),
    ("FR-21", "Match notifications", "Email alerts for interest and new messages.", "Sprint 2", "notifications/tasks.py, notifications/email.py"),
    ("FR-22", "Responsive web UI", "Core flows work on desktop and mobile.", "Sprint 1–2", "templates/base.html, static/css/"),
    ("FR-23", "Input validation", "Field-level errors for invalid or edge-case input.", "Sprint 1", "*/forms.py, */serializers.py"),
    ("FR-24", "Listing detail view", "Full listing page with author and actions.", "Sprint 1", "listings/views.py, templates/listings/detail.html"),
    ("FR-25", "Keyword search", "Search listings by title/description text.", "Sprint 2", "listings/filters.py, templates/listings/search.html"),
    ("FR-26", "In-app notifications", "Unread alerts for matches, messages, ratings.", "Sprint 2", "notifications/models.py, templates/notifications/"),
    ("FR-27", "Personal dashboard", "Member view of listings, matches, chats.", "Sprint 2", "accounts/views.py, templates/accounts/dashboard.html"),
    ("FR-28", "Moderator report workflow", "Review, resolve, or dismiss user reports.", "Sprint 2", "moderation/views.py, moderation/models.py"),
]

NON_FUNCTIONAL_REQUIREMENTS = [
    ("NFR-01", "Security — authentication", "Hashed passwords; protected routes require auth.", "Sprint 1", "config/settings.py, accounts/authentication.py"),
    ("NFR-02", "Security — CSRF and XSS", "CSRF tokens and template auto-escaping.", "Sprint 1", "Django middleware, DRF settings"),
    ("NFR-03", "Security — authorisation", "Object-level permissions by role.", "Sprint 1", "accounts/permissions.py, */permissions.py"),
    ("NFR-04", "Data privacy", "Suburb/postcode only; no exact addresses stored.", "Sprint 1", "profiles/models.py, profiles/forms.py"),
    ("NFR-05", "Performance — page response", "Listing pages load within 3 seconds.", "Sprint 2", "DB indexes, pagination, query optimisation"),
    ("NFR-06", "Performance — real-time chat", "Messages delivered within 2 seconds.", "Sprint 2", "chat/consumers.py, Redis CHANNEL_LAYERS"),
    ("NFR-07", "Reliability — data persistence", "PostgreSQL with transactional integrity.", "Sprint 1", "config/settings.py, migrations"),
    ("NFR-08", "Scalability", "Stateless REST + ASGI for horizontal scaling.", "Sprint 2", "config/asgi.py, docs/DEPLOYMENT.md"),
    ("NFR-09", "Usability", "Core journeys ≤ 5 clicks with clear errors.", "Sprint 1–2", "Templates, Sprint review UX checks"),
    ("NFR-10", "Accessibility", "WCAG 2.1 Level A basics (semantic HTML, contrast).", "Sprint 2", "templates/base.html, docs/ACCESSIBILITY.md"),
    ("NFR-11", "Maintainability", "Modular Django apps per domain.", "Sprint 1", "backend/ app structure, docs/ARCHITECTURE.md"),
    ("NFR-12", "Testability", "≥ 30 automated tests (pytest-django).", "Sprint 1–2", "tests/, conftest.py, pytest.ini"),
    ("NFR-13", "Observability", "Structured logging for errors and moderation.", "Sprint 2", "config/settings.py (LOGGING)"),
    ("NFR-14", "Deployability", "Docker Compose + env-based staging deploy.", "Sprint 2", "docker-compose.yml, docs/DEPLOYMENT.md"),
    ("NFR-15", "Email deliverability", "Configurable SMTP / console backend.", "Sprint 1–2", "config/settings.py, notifications/email.py"),
    ("NFR-16", "Reliability — availability", "≥ 99% uptime; friendly error pages.", "Sprint 2", "templates/500.html, health-check endpoint"),
    ("NFR-17", "Portability", "Runs in Docker/staging via env vars only.", "Sprint 1–2", "docker-compose.yml, .env.example, settings.py"),
    ("NFR-18", "Security — rate limiting", "Throttle login, report, and chat endpoints.", "Sprint 2", "config/settings.py, accounts/throttling.py"),
]

EXTENDED_REQUIREMENTS = [
    ("EXT-01", "Logging", "Record errors and key events for audit/debug.", "Sprint 2", "config/settings.py (LOGGING), views/consumers"),
    ("EXT-02", "Monitoring and alerting", "Health-check endpoint; deployment monitoring notes.", "Sprint 2", "config/views.py, docs/DEPLOYMENT.md"),
    ("EXT-03", "Usage analytics", "Admin summary of registrations, listings, matches.", "Sprint 2", "analytics/models.py, analytics/views.py"),
    ("EXT-04", "Backup and disaster recovery", "DB backup script and restore documentation.", "Sprint 2", "scripts/backup_db.sh, docs/DEPLOYMENT.md"),
    ("EXT-05", "Rate limiting (API)", "Request quotas on public/authenticated APIs.", "Sprint 2", "DRF throttle classes in settings.py"),
]


def _arc_arrow(ax, cx, cy, r, theta1, theta2, color, lw=10):
    ax.add_patch(plt.matplotlib.patches.Arc(
        (cx, cy), 2 * r, 2 * r, angle=0,
        theta1=theta1, theta2=theta2,
        color=color, linewidth=lw, zorder=3,
    ))
    rad = math.radians(theta2)
    tip_x = cx + r * math.cos(rad)
    tip_y = cy + r * math.sin(rad)
    dx = -math.sin(rad) * 0.18
    dy = math.cos(rad) * 0.18
    ax.annotate(
        "", xy=(tip_x, tip_y), xytext=(tip_x - dx, tip_y - dy),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=lw * 0.7, mutation_scale=14),
        zorder=4,
    )


def _label_on_arc(ax, cx, cy, r, angle_deg, text, color="#333333", fontsize=8, offset=0.35):
    rad = math.radians(angle_deg)
    x = cx + (r + offset) * math.cos(rad)
    y = cy + (r + offset) * math.sin(rad)
    rotation = angle_deg + (180 if 90 < angle_deg < 270 else 0)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize,
            fontweight="bold", color=color, rotation=rotation, zorder=6)


def _draw_gear(ax, cx, cy, r):
    ax.add_patch(plt.Circle((cx, cy), r * 1.15, facecolor=GEAR_COLOR, edgecolor="#D5C9BA",
                             linewidth=1, alpha=0.55, zorder=1))
    for angle in np.linspace(0, 360, 16, endpoint=False):
        rad = math.radians(angle)
        ax.plot(
            [cx + r * 1.05 * math.cos(rad), cx + r * 1.22 * math.cos(rad)],
            [cy + r * 1.05 * math.sin(rad), cy + r * 1.22 * math.sin(rad)],
            color="#D5C9BA", linewidth=2, alpha=0.6, zorder=2,
        )


def _draw_sprint(ax, cx, cy, r, sprint_num, show_plan_left=True, show_launch_right=True):
    _draw_gear(ax, cx, cy, r)
    _arc_arrow(ax, cx, cy, r, 315, 45, DARK_BLUE)
    _arc_arrow(ax, cx, cy, r, 45, 135, LIGHT_BLUE)
    _arc_arrow(ax, cx, cy, r, 135, 225, YELLOW)
    _arc_arrow(ax, cx, cy, r, 225, 285, YELLOW)
    _label_on_arc(ax, cx, cy, r, 0, "DEVELOP", DARK_BLUE)
    _label_on_arc(ax, cx, cy, r, 90, "TEST", LIGHT_BLUE)
    _label_on_arc(ax, cx, cy, r, 180, "DEPLOY", "#B8860B")
    _label_on_arc(ax, cx, cy, r, 255, "REVIEW", "#B8860B")
    ax.text(cx, cy, f"SPRINT {sprint_num}", ha="center", va="center",
            fontsize=13, fontweight="bold", color=SPRINT_COLORS[sprint_num], zorder=6)
    if show_plan_left:
        ax.annotate("", xy=(cx - r * 0.85, cy - r * 0.55), xytext=(cx - r * 2.0, cy - r * 0.55),
                    arrowprops=dict(arrowstyle="-|>", color=DARK_BLUE, lw=3, mutation_scale=16), zorder=5)
        ax.text(cx - r * 1.55, cy - r * 0.55 - 0.28, "PLAN",
                ha="center", fontsize=9, fontweight="bold", color=DARK_BLUE, zorder=6)
    ax.add_patch(plt.matplotlib.patches.Arc(
        (cx, cy), 2 * r * 0.72, 2 * r * 0.45, angle=0,
        theta1=200, theta2=340, color=RED, linewidth=9, zorder=3,
    ))
    ax.text(cx, cy - r * 0.78, "DESIGN", ha="center", fontsize=9, fontweight="bold", color=RED, zorder=6)
    if show_launch_right:
        ax.annotate("", xy=(cx + r * 2.0, cy - r * 0.55), xytext=(cx + r * 0.9, cy - r * 0.55),
                    arrowprops=dict(arrowstyle="-|>", color=RED, lw=3, mutation_scale=16), zorder=5)
        ax.text(cx + r * 1.55, cy - r * 0.55 - 0.28, "LAUNCH",
                ha="center", fontsize=9, fontweight="bold", color=RED, zorder=6)


def create_agile_diagram():
    fig, ax = plt.subplots(figsize=(12, 5.5), dpi=150)
    ax.set_xlim(-0.5, 12.5)
    ax.set_ylim(-0.5, 5.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")
    ax.text(6, 5.0, "AGILE METHODOLOGY", ha="center", va="center",
            fontsize=20, fontweight="bold", color="#333333", zorder=6)
    ax.text(6, 4.55, "Local Community Skill Exchange and Help Board  |  Group F",
            ha="center", va="center", fontsize=9, color="#666666", zorder=6)
    r, cy, cx1, cx2 = 1.55, 2.3, 3.2, 9.2
    _draw_sprint(ax, cx1, cy, r, 1, show_plan_left=True, show_launch_right=True)
    _draw_sprint(ax, cx2, cy, r, 2, show_plan_left=False, show_launch_right=True)
    ax.annotate("", xy=(cx2 - r * 2.0, cy - r * 0.55), xytext=(cx1 + r * 2.0, cy - r * 0.55),
                arrowprops=dict(arrowstyle="-|>", color=RED, lw=2.5, mutation_scale=14), zorder=5)
    ax.text(11.5, cy - r * 0.55, "Prototype\nDelivered", ha="center", va="center",
            fontsize=8, fontweight="bold", color=RED, zorder=6)
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight", facecolor="white", pad_inches=0.2)
    plt.close(fig)
    buf.seek(0)
    return buf


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return table


def build_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    h = doc.add_heading("Assessment 2 — Project Documentation", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local Community Skill Exchange and Help Board")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed")
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "This document combines the agile project plan and requirements specification for a "
        "community platform where neighbors post help requests and offers, match by location, "
        "chat in real time, and rate each other. The project runs in 2 sprints of 2 weeks each."
    )

    # --- Agile Plan ---
    doc.add_heading("Agile Project Plan", level=2)
    doc.add_heading("Agile Sprint Diagram", level=3)
    doc.add_picture(create_agile_diagram(), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Sprint Summary", level=3)
    add_table(doc,
        ["Sprint", "Duration", "Objective", "Outcome"],
        [
            ["Sprint 1", "Week 1–2", "User accounts, profiles, listings, admin",
             "Register, profile, post/browse listings; 12 unit tests"],
            ["Sprint 2", "Week 3–4", "Search, chat, ratings, safety, deployment",
             "Full prototype with chat, ratings, moderation; 18 integration tests"],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("Sprint 1 Backlog", level=3)
    add_table(doc, ["Task", "Requirements", "Implementation"],
        [
            ["Django + PostgreSQL setup", "NFR-07, NFR-11", "backend/config/, docker-compose.yml"],
            ["Registration & email verification", "FR-01, FR-02", "backend/accounts/"],
            ["Login, logout, password reset", "FR-03, FR-04", "backend/accounts/"],
            ["User profiles", "FR-05", "backend/profiles/"],
            ["Roles and permissions", "FR-06", "backend/accounts/permissions.py"],
            ["Listings and categories", "FR-07 – FR-10", "backend/listings/"],
            ["Input validation & listing detail", "FR-23, FR-24", "backend/listings/, */forms.py"],
            ["Admin moderation panel", "FR-11", "backend/moderation/"],
            ["Docker env configuration", "NFR-17", "docker-compose.yml, .env.example"],
            ["Unit tests", "NFR-12", "backend/tests/ (12 tests)"],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("Sprint 2 Backlog", level=3)
    add_table(doc, ["Task", "Requirements", "Implementation"],
        [
            ["Nearby, keyword search, dashboard", "FR-12, FR-25, FR-27", "backend/listings/filters.py"],
            ["Match, chat, in-app notifications", "FR-14 – FR-16, FR-26", "backend/chat/, notifications/"],
            ["Ratings and reviews", "FR-17, FR-18", "backend/reviews/"],
            ["Report, block, moderator workflow", "FR-19, FR-20, FR-28", "backend/moderation/"],
            ["Rate limiting & logging", "NFR-18, EXT-01, EXT-05", "settings.py, throttling.py"],
            ["Health check, backup, deploy", "EXT-02 – EXT-04, NFR-16", "docs/DEPLOYMENT.md, scripts/"],
            ["Integration tests & deploy", "NFR-12 – NFR-14", "backend/tests/, docs/DEPLOYMENT.md"],
        ],
    )

    doc.add_page_break()

    # --- Functional Requirements ---
    doc.add_heading("Functional Requirements", level=2)
    doc.add_paragraph(
        "Each functional requirement describes what the system must do. "
        "The Sprint column indicates when the requirement is implemented; "
        "Implementation Location shows the planned module path under backend/."
    )
    doc.add_paragraph()
    add_table(doc,
        ["ID", "Requirement", "Description", "Sprint", "Implementation Location"],
        [(rid, name, desc, sprint, loc) for rid, name, desc, sprint, loc in FUNCTIONAL_REQUIREMENTS],
        font_size=8,
    )

    doc.add_page_break()

    # --- Non-Functional Requirements ---
    doc.add_heading("Non-Functional Requirements", level=2)
    doc.add_paragraph(
        "Non-functional requirements define quality attributes: security, performance, "
        "usability, testability, and deployability."
    )
    doc.add_paragraph()
    add_table(doc,
        ["ID", "Requirement", "Description", "Sprint", "Implementation Location"],
        [(rid, name, desc, sprint, loc) for rid, name, desc, sprint, loc in NON_FUNCTIONAL_REQUIREMENTS],
        font_size=8,
    )

    doc.add_paragraph()
    doc.add_heading("Extended Requirements", level=2)
    doc.add_paragraph(
        "Extended requirements (GeeksforGeeks) support monitoring, reliability, and future "
        "expansion beyond core features: logging, monitoring, analytics, backup, and rate limiting."
    )
    doc.add_paragraph()
    add_table(doc,
        ["ID", "Requirement", "Description", "Sprint", "Implementation Location"],
        [(rid, name, desc, sprint, loc) for rid, name, desc, sprint, loc in EXTENDED_REQUIREMENTS],
        font_size=8,
    )

    doc.add_paragraph()
    doc.add_heading("Requirements Traceability", level=3)
    add_table(doc, ["Sprint", "Functional", "Non-Functional", "Extended"],
        [
            ["Sprint 1", "FR-01 – FR-11, FR-22 (base), FR-23 – FR-24",
             "NFR-01 – NFR-04, NFR-07, NFR-11, NFR-12 (12 tests), NFR-15, NFR-17", "—"],
            ["Sprint 2", "FR-12 – FR-28, FR-22 (complete)",
             "NFR-05 – NFR-06, NFR-08 – NFR-18", "EXT-01 – EXT-05"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Acceptance: prototype complete when users can register with validated input, "
        "post/browse/detail listings, keyword and nearby search, chat in real time, "
        "receive in-app notifications, leave ratings, and admins/moderators resolve reports — "
        "with ≥ 30 passing automated tests and documented backup/health-check procedures."
    )

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
