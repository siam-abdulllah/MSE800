#!/usr/bin/env python3
"""Generate standalone Functional & Non-Functional Requirements Word document."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "Functional_and_NonFunctional_Requirements.docx",
)

GEEKSFORGEEKS_URL = (
    "https://www.geeksforgeeks.org/software-engineering/"
    "functional-vs-non-functional-requirements/"
)

FUNCTIONAL_SECTIONS = [
    ("Authentication and User Management", [
        ("FR-01", "User registration", "Create account with email and password.", "Sprint 1", "accounts/views.py, accounts/models.py"),
        ("FR-02", "Email verification", "Verification link; unverified users restricted.", "Sprint 1", "accounts/views.py, notifications/email.py"),
        ("FR-03", "Login and logout", "Secure session authentication.", "Sprint 1", "accounts/views.py, accounts/urls.py"),
        ("FR-04", "Password reset", "Time-limited reset via email.", "Sprint 1", "accounts/views.py"),
        ("FR-05", "User profile", "Skills, bio, suburb, optional photo.", "Sprint 1", "profiles/models.py, profiles/views.py"),
        ("FR-06", "Role-based access", "Member, Moderator, Admin roles.", "Sprint 1", "accounts/permissions.py"),
        ("FR-23", "Input validation", "Clear field-level errors for invalid input.", "Sprint 1", "*/forms.py, */serializers.py"),
    ]),
    ("Listings and Discovery", [
        ("FR-07", "Create help listing", "Post need/can help with category and suburb.", "Sprint 1", "listings/models.py, listings/views.py"),
        ("FR-08", "Browse listings", "Public feed of active listings.", "Sprint 1", "listings/views.py, templates/listings/"),
        ("FR-09", "Listing categories", "Predefined help categories.", "Sprint 1", "listings/models.py (Category)"),
        ("FR-10", "Edit and close listing", "Owners edit or close listings.", "Sprint 1", "listings/views.py"),
        ("FR-24", "Listing detail view", "Full listing page with actions.", "Sprint 1", "listings/views.py, templates/listings/detail.html"),
        ("FR-12", "Nearby search", "Filter by suburb, postcode, or radius.", "Sprint 2", "listings/filters.py, listings/services/location.py"),
        ("FR-13", "Skill-based matching", "Filter by skill or category.", "Sprint 2", "listings/filters.py"),
        ("FR-25", "Keyword search", "Search title and description text.", "Sprint 2", "listings/filters.py, templates/base.html"),
    ]),
    ("Matching, Communication, and Trust", [
        ("FR-14", "Express interest / match", "Open private chat between users.", "Sprint 2", "listings/views.py, chat/models.py"),
        ("FR-15", "Real-time chat", "Live messaging via Channels + Redis.", "Sprint 2", "chat/consumers.py, config/asgi.py"),
        ("FR-16", "Chat history", "Persisted scrollable messages.", "Sprint 2", "chat/models.py, chat/views.py"),
        ("FR-17", "Star ratings", "1–5 stars with optional review.", "Sprint 2", "reviews/models.py, reviews/views.py"),
        ("FR-18", "Aggregate rating display", "Average rating on profiles.", "Sprint 2", "profiles/views.py, reviews/models.py"),
        ("FR-21", "Email notifications", "Alerts for matches and messages.", "Sprint 2", "notifications/tasks.py"),
        ("FR-26", "In-app notification centre", "Unread badge and notification inbox.", "Sprint 2", "notifications/models.py"),
        ("FR-27", "Personal dashboard", "Overview of listings, matches, chats.", "Sprint 2", "accounts/views.py (DashboardView)"),
    ]),
    ("Safety, Moderation, and Presentation", [
        ("FR-11", "Admin moderation panel", "Manage users, listings, reports.", "Sprint 1", "moderation/admin.py"),
        ("FR-19", "Report user or listing", "Reports enter moderation queue.", "Sprint 2", "moderation/models.py, moderation/views.py"),
        ("FR-20", "Block user", "Block messaging and hide listings.", "Sprint 2", "moderation/models.py, chat/services.py"),
        ("FR-28", "Moderator report workflow", "Resolve or dismiss reports with notes.", "Sprint 2", "moderation/views.py"),
        ("FR-22", "Responsive web UI", "Desktop and mobile browser support.", "Sprint 1–2", "templates/base.html, static/css/"),
    ]),
]

NFR_SECTIONS = [
    ("Performance", [
        ("NFR-05", "Page response time", "Listing pages load within 3 seconds.", "Sprint 2", "DB indexes, pagination"),
        ("NFR-06", "Real-time chat latency", "Messages delivered within 2 seconds.", "Sprint 2", "chat/consumers.py, Redis"),
    ]),
    ("Security", [
        ("NFR-01", "Authentication", "Hashed passwords; protected routes.", "Sprint 1", "config/settings.py, accounts/authentication.py"),
        ("NFR-02", "CSRF and XSS protection", "CSRF tokens; template escaping.", "Sprint 1", "Django middleware"),
        ("NFR-03", "Authorisation", "Role and object-level permissions.", "Sprint 1", "*/permissions.py"),
        ("NFR-04", "Data privacy", "Suburb/postcode only; no street addresses.", "Sprint 1", "profiles/models.py"),
        ("NFR-18", "Rate limiting", "Throttle auth, report, chat endpoints.", "Sprint 2", "config/settings.py, throttling.py"),
    ]),
    ("Usability", [
        ("NFR-09", "Ease of use", "Core journeys ≤ 5 clicks.", "Sprint 1–2", "Templates, wireframes"),
        ("NFR-10", "Accessibility", "WCAG 2.1 Level A basics.", "Sprint 2", "templates/base.html"),
    ]),
    ("Reliability", [
        ("NFR-07", "Data persistence", "PostgreSQL with ACID transactions.", "Sprint 1", "config/settings.py, migrations"),
        ("NFR-16", "Availability", "≥ 99% uptime; friendly error pages.", "Sprint 2", "templates/500.html, health-check"),
    ]),
    ("Scalability", [
        ("NFR-08", "Horizontal scaling", "Stateless REST + ASGI WebSockets.", "Sprint 2", "config/asgi.py"),
    ]),
    ("Maintainability", [
        ("NFR-11", "Modular architecture", "Separate Django apps per domain.", "Sprint 1", "backend/ layout"),
        ("NFR-12", "Automated testing", "≥ 30 pytest-django tests.", "Sprint 1–2", "tests/, pytest.ini"),
        ("NFR-13", "Logging", "Errors and key events logged.", "Sprint 2", "config/settings.py (LOGGING)"),
        ("NFR-14", "Deployability", "Docker Compose staging deploy.", "Sprint 2", "docker-compose.yml"),
        ("NFR-15", "Email deliverability", "Configurable SMTP / console backend.", "Sprint 1–2", "notifications/email.py"),
    ]),
    ("Portability", [
        ("NFR-17", "Cross-environment execution", "Env-var config for local/Docker/staging.", "Sprint 1–2", ".env.example, settings.py"),
    ]),
]

EXTENDED_REQUIREMENTS = [
    ("EXT-01", "Logging", "Audit log for failed logins, reports, chat errors.", "Sprint 2", "config/settings.py (LOGGING)"),
    ("EXT-02", "Monitoring and alerting", "Health-check for DB and Redis.", "Sprint 2", "config/views.py, DEPLOYMENT.md"),
    ("EXT-03", "Usage analytics", "Admin summary of platform activity.", "Sprint 2", "analytics/models.py"),
    ("EXT-04", "Backup and disaster recovery", "DB backup script and restore docs.", "Sprint 2", "scripts/backup_db.sh"),
    ("EXT-05", "API rate limiting", "Request quotas on APIs.", "Sprint 2", "DRF throttle classes"),
]


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_size)
    return table


def build_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    title = doc.add_heading("Functional and Non-Functional Requirements", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local Community Skill Exchange and Help Board")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed  |  MSE800 Assessment 2")
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "This document follows the requirements classification from GeeksforGeeks: "
        "functional requirements define what the system must do; non-functional requirements "
        "define how it must perform; extended requirements cover logging, monitoring, "
        "analytics, backup, and rate limiting."
    )
    doc.add_paragraph(f"Reference: {GEEKSFORGEEKS_URL}")

    doc.add_heading("Functional vs Non-Functional — Summary", level=2)
    add_table(doc,
        ["Functional Requirements", "Non-Functional Requirements"],
        [
            ["Define what the system should do", "Define how the system should perform"],
            ["Features: login, listings, chat, ratings", "Quality: speed, security, usability, reliability"],
            ["Directly visible to users", "Validated with metrics and SLAs"],
            ["Documented as user stories / use cases", "Documented as technical quality criteria"],
        ],
        font_size=10,
    )

    doc.add_page_break()
    doc.add_heading("Functional Requirements", level=2)
    doc.add_paragraph(
        "What the system must do — observable features and operations (GeeksforGeeks)."
    )
    headers = ["ID", "Requirement", "Description", "Sprint", "Implementation Location"]
    for section_name, rows in FUNCTIONAL_SECTIONS:
        doc.add_heading(section_name, level=3)
        add_table(doc, headers, rows, font_size=8)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Non-Functional Requirements", level=2)
    doc.add_paragraph(
        "How the system must perform — grouped by performance, security, usability, "
        "reliability, scalability, maintainability, and portability (GeeksforGeeks)."
    )
    for section_name, rows in NFR_SECTIONS:
        doc.add_heading(section_name, level=3)
        add_table(doc, headers, rows, font_size=8)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Extended Requirements", level=2)
    doc.add_paragraph(
        "Additional capabilities beyond core features: logging, monitoring, analytics, "
        "backup, and API protection (GeeksforGeeks extended requirements)."
    )
    add_table(doc, headers, EXTENDED_REQUIREMENTS, font_size=8)

    doc.add_paragraph()
    doc.add_heading("Sprint Traceability", level=2)
    add_table(doc, ["Sprint", "Functional", "Non-Functional", "Extended"],
        [
            ["Sprint 1 (Weeks 1–2)", "FR-01 – FR-11, FR-22 (base), FR-23 – FR-24",
             "NFR-01 – NFR-04, NFR-07, NFR-11, NFR-12 (12 tests), NFR-15, NFR-17", "—"],
            ["Sprint 2 (Weeks 3–4)", "FR-12 – FR-28, FR-22 (complete)",
             "NFR-05 – NFR-06, NFR-08 – NFR-18", "EXT-01 – EXT-05"],
        ],
        font_size=9,
    )

    doc.add_paragraph()
    doc.add_heading("Project Example — Help Board Scenario", level=2)
    doc.add_paragraph(
        "Functional: users log in, post and browse listings, search nearby, chat after "
        "matching, and receive notifications. "
        "Non-functional: pages respond in under 3 seconds, sessions are protected, suburb-only "
        "privacy, ≥ 99% demo availability, Docker portability. "
        "Extended: failed logins logged, database backups documented, auth endpoints rate-limited."
    )

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
