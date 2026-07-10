#!/usr/bin/env python3
"""Generate a simple Agile Project Management Plan with diagram."""

import io
import os
import math
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Arc
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "Local_Community_Skill_Exchange_Agile_Plan.docx",
)

DARK_BLUE = "#1B3A6B"
LIGHT_BLUE = "#5BA3E8"
YELLOW = "#E8B923"
RED = "#C0392B"
GEAR_COLOR = "#E8DDD0"
SPRINT_COLORS = {1: DARK_BLUE, 2: "#D4A017"}


def _arc_arrow(ax, cx, cy, r, theta1, theta2, color, lw=10):
    ax.add_patch(Arc(
        (cx, cy), 2 * r, 2 * r, angle=0,
        theta1=theta1, theta2=theta2,
        color=color, linewidth=lw, zorder=3,
    ))
    rad = math.radians(theta2)
    tip_x = cx + r * math.cos(rad)
    tip_y = cy + r * math.sin(rad)
    dx = -math.sin(rad) * 0.18
    dy = math.cos(rad) * 0.18
    ax.annotate(
        "", xy=(tip_x, tip_y), xytext=(tip_x - dx, tip_y - dy),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=lw * 0.7, mutation_scale=14),
        zorder=4,
    )


def _label_on_arc(ax, cx, cy, r, angle_deg, text, color="#333333", fontsize=8, offset=0.35):
    rad = math.radians(angle_deg)
    x = cx + (r + offset) * math.cos(rad)
    y = cy + (r + offset) * math.sin(rad)
    rotation = angle_deg + (180 if 90 < angle_deg < 270 else 0)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize,
            fontweight="bold", color=color, rotation=rotation, zorder=6)


def _draw_gear(ax, cx, cy, r):
    ax.add_patch(Circle((cx, cy), r * 1.15, facecolor=GEAR_COLOR, edgecolor="#D5C9BA",
                        linewidth=1, alpha=0.55, zorder=1))
    for angle in np.linspace(0, 360, 16, endpoint=False):
        rad = math.radians(angle)
        ax.plot(
            [cx + r * 1.05 * math.cos(rad), cx + r * 1.22 * math.cos(rad)],
            [cy + r * 1.05 * math.sin(rad), cy + r * 1.22 * math.sin(rad)],
            color="#D5C9BA", linewidth=2, alpha=0.6, zorder=2,
        )


def _draw_sprint(ax, cx, cy, r, sprint_num, show_plan_left=True, show_launch_right=True):
    _draw_gear(ax, cx, cy, r)

    _arc_arrow(ax, cx, cy, r, 315, 45, DARK_BLUE)
    _arc_arrow(ax, cx, cy, r, 45, 135, LIGHT_BLUE)
    _arc_arrow(ax, cx, cy, r, 135, 225, YELLOW)
    _arc_arrow(ax, cx, cy, r, 225, 285, YELLOW)

    _label_on_arc(ax, cx, cy, r, 0, "DEVELOP", DARK_BLUE)
    _label_on_arc(ax, cx, cy, r, 90, "TEST", LIGHT_BLUE)
    _label_on_arc(ax, cx, cy, r, 180, "DEPLOY", "#B8860B")
    _label_on_arc(ax, cx, cy, r, 255, "REVIEW", "#B8860B")

    ax.text(cx, cy, f"SPRINT {sprint_num}", ha="center", va="center",
            fontsize=13, fontweight="bold", color=SPRINT_COLORS[sprint_num], zorder=6)

    if show_plan_left:
        ax.annotate(
            "", xy=(cx - r * 0.85, cy - r * 0.55),
            xytext=(cx - r * 2.0, cy - r * 0.55),
            arrowprops=dict(arrowstyle="-|>", color=DARK_BLUE, lw=3, mutation_scale=16),
            zorder=5,
        )
        ax.text(cx - r * 1.55, cy - r * 0.55 - 0.28, "PLAN",
                ha="center", fontsize=9, fontweight="bold", color=DARK_BLUE, zorder=6)

    ax.add_patch(Arc(
        (cx, cy), 2 * r * 0.72, 2 * r * 0.45, angle=0,
        theta1=200, theta2=340, color=RED, linewidth=9, zorder=3,
    ))
    ax.text(cx, cy - r * 0.78, "DESIGN", ha="center", fontsize=9,
            fontweight="bold", color=RED, zorder=6)

    if show_launch_right:
        ax.annotate(
            "", xy=(cx + r * 2.0, cy - r * 0.55),
            xytext=(cx + r * 0.9, cy - r * 0.55),
            arrowprops=dict(arrowstyle="-|>", color=RED, lw=3, mutation_scale=16),
            zorder=5,
        )
        ax.text(cx + r * 1.55, cy - r * 0.55 - 0.28, "LAUNCH",
                ha="center", fontsize=9, fontweight="bold", color=RED, zorder=6)


def create_agile_diagram():
    """Agile methodology diagram — circular sprint cycles like reference template."""
    fig, ax = plt.subplots(figsize=(12, 5.5), dpi=150)
    ax.set_xlim(-0.5, 12.5)
    ax.set_ylim(-0.5, 5.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(6, 5.0, "AGILE METHODOLOGY", ha="center", va="center",
            fontsize=20, fontweight="bold", color="#333333", zorder=6)
    ax.text(6, 4.55,
            "Local Community Skill Exchange and Help Board  |  Group F",
            ha="center", va="center", fontsize=9, color="#666666", zorder=6)

    r = 1.55
    cy = 2.3
    cx1, cx2 = 3.2, 9.2

    _draw_sprint(ax, cx1, cy, r, 1, show_plan_left=True, show_launch_right=True)
    _draw_sprint(ax, cx2, cy, r, 2, show_plan_left=False, show_launch_right=True)

    ax.annotate(
        "", xy=(cx2 - r * 2.0, cy - r * 0.55),
        xytext=(cx1 + r * 2.0, cy - r * 0.55),
        arrowprops=dict(arrowstyle="-|>", color=RED, lw=2.5, mutation_scale=14),
        zorder=5,
    )

    ax.text(11.5, cy - r * 0.55, "Prototype\nDelivered", ha="center", va="center",
            fontsize=8, fontweight="bold", color=RED, zorder=6)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight", facecolor="white", pad_inches=0.2)
    plt.close(fig)
    buf.seek(0)
    return buf


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    return table


def build_document():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    h = doc.add_heading("Agile Project Management Plan", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local Community Skill Exchange and Help Board")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed")
    r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "Agile (Scrum) plan for building a community platform where neighbors post help "
        "requests and offers, match by location, chat, and rate each other. "
        "The project runs in 2 sprints of 2 weeks each."
    )

    doc.add_paragraph()

    doc.add_heading("Agile Sprint Diagram", level=2)
    doc.add_picture(create_agile_diagram(), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("Sprint Summary", level=2)
    add_table(doc,
        ["Sprint", "Duration", "Objective", "Outcome"],
        [
            ["Sprint 1", "Week 1–2",
             "Set up project and build user accounts & listings",
             "Users can register, create profiles, and post/browse help listings"],
            ["Sprint 2", "Week 3–4",
             "Add search, chat, ratings, and deploy",
             "Full working prototype with chat, ratings, and safety features"],
        ],
    )

    doc.add_paragraph()

    doc.add_heading("Sprint 1 — User Accounts & Listings", level=2)
    add_table(doc,
        ["Tasks", "Deliverables"],
        [
            ["Set up Django + PostgreSQL", "Working dev environment"],
            ["User registration & email verification", "Members can sign up and log in"],
            ["Profile page (skills, suburb, photo)", "User profiles with basic info"],
            ["Create and browse help listings", "Post 'I need help' / 'I can help' listings"],
            ["Admin moderation panel", "Admin can manage users and posts"],
            ["Unit tests (pytest-django)", "12 tests passing"],
        ],
    )

    doc.add_paragraph()

    doc.add_heading("Sprint 2 — Matching, Chat & Ratings", level=2)
    add_table(doc,
        ["Tasks", "Deliverables"],
        [
            ["Nearby search by suburb/postcode", "Location-based listing filter"],
            ["Real-time chat (Django Channels + Redis)", "Matched users can chat live"],
            ["Rating and review system", "Star ratings shown on profiles"],
            ["Report and block user features", "Safety controls for the community"],
            ["Email match notifications", "Users notified of new matches"],
            ["Integration testing & deployment", "18 tests passing, prototype live"],
        ],
    )

    doc.add_paragraph()

    doc.add_heading("Activity Log & Progress Updates", level=2)
    add_table(doc,
        ["Week", "Sprint", "Activity", "Status"],
        [
            ["Week 1", "Sprint 1", "Sprint planning, project setup, user auth", "Done"],
            ["Week 2", "Sprint 1", "Profiles, listings, admin panel, unit tests", "Done"],
            ["Week 2", "Sprint 1", "Sprint review & retrospective", "Done"],
            ["Week 3", "Sprint 2", "Sprint planning, nearby search, chat setup", "Done"],
            ["Week 4", "Sprint 2", "Ratings, report/block, testing, deployment", "Done"],
            ["Week 4", "Sprint 2", "Sprint review & retrospective", "Done"],
        ],
    )

    doc.add_paragraph()

    doc.add_heading("Sprint Reviews & Retrospectives", level=2)
    add_table(doc,
        ["Sprint", "Review (Demo)", "Retrospective"],
        [
            ["Sprint 1",
             "Demo: register user, create profile, post listing, browse home page",
             "Went well: clear task split, on-time delivery. Improve: create wireframes earlier"],
            ["Sprint 2",
             "Demo: search nearby, chat live, leave rating, report user, deploy prototype",
             "Went well: chat worked first try, all tests pass. Improve: allow more time for deployment"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Final Outcome: A working prototype where users register, post and browse help listings, "
        "search nearby, chat in real time, leave ratings, and admins moderate content — "
        "ready for software engineering assessment."
    )

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
