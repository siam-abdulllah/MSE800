#!/usr/bin/env python3
"""Generate a simple Waterfall Project Management Word document."""

import io
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "Local_Community_Skill_Exchange_Waterfall_Diagram.docx",
)

BOX_COLOR = "#B8D4F0"
BORDER_COLOR = "#2C5F8A"


def create_waterfall_diagram():
    """Simple vertical waterfall diagram matching the template style."""
    fig, ax = plt.subplots(figsize=(6, 10), dpi=150)
    ax.set_xlim(0, 6)
    ax.set_ylim(0, 10)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    phases = [
        "Requirements Analysis",
        "System Design",
        "Development",
        "Testing",
        "Deployment",
        "Maintenance",
    ]
    widths = [4.2, 3.8, 4.0, 3.4, 3.6, 3.4]
    y_positions = [8.4, 6.9, 5.4, 3.9, 2.4, 0.9]

    for i, (label, y, w) in enumerate(zip(phases, y_positions, widths)):
        x = (6 - w) / 2
        box = FancyBboxPatch(
            (x, y), w, 0.9,
            boxstyle="round,pad=0.04,rounding_size=0.1",
            facecolor=BOX_COLOR,
            edgecolor=BORDER_COLOR,
            linewidth=1.2,
        )
        ax.add_patch(box)
        ax.text(3, y + 0.45, label, ha="center", va="center",
                fontsize=11, fontweight="bold", color="#1A2E44")

        if i < len(phases) - 1:
            ax.add_patch(FancyArrowPatch(
                (3, y), (3, y_positions[i + 1] + 0.9),
                arrowstyle="-|>", mutation_scale=15,
                color=BORDER_COLOR, linewidth=2,
                shrinkA=3, shrinkB=3,
            ))

    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Waterfall Methodology", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Local Community Skill Exchange and Help Board")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Group F  |  Syed Abdullah Maaz  |  Fahad Ahmed")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Brief intro
    doc.add_paragraph(
        "This diagram shows the Waterfall Model for our community platform where "
        "neighbors post help requests and offers, match by location and skill, "
        "and rate each other. Each phase is completed before moving to the next."
    )

    doc.add_paragraph()

    # Diagram
    doc.add_picture(create_waterfall_diagram(), width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Single phase table
    doc.add_heading("Project Phases", level=2)

    phases = [
        ("Requirements Analysis", "Week 1–2", "2 weeks",
         "Define what the platform needs: user registration, help posts, nearby search, chat, and ratings.",
         "Focus on core features first — registration, listings, and messaging."),
        ("System Design", "Week 3–4", "2 weeks",
         "Plan the architecture using Django, PostgreSQL, GeoDjango, and Django Channels.",
         "Create database schema, API design, and UI wireframes."),
        ("Development", "Week 5–9", "5 weeks",
         "Build the prototype: backend APIs, frontend UI, real-time chat, and user profiles.",
         "Backend first, then connect the frontend."),
        ("Testing", "Week 10–11", "2 weeks",
         "Test all features using pytest-django. Fix bugs and check usability.",
         "Test registration, search, chat, and rating workflows."),
        ("Deployment", "Week 12", "1 week",
         "Deploy the working prototype to a server. Set up database, Redis, and email.",
         "Run final checks before going live."),
        ("Maintenance", "Week 13+", "Ongoing",
         "Fix bugs, update content, and manage the moderation panel.",
         "Monitor performance and respond to user feedback."),
    ]

    table = doc.add_table(rows=len(phases) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Timeline", "Duration", "Description", "Notes"]

    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        shade_cell(cell, "2C5F8A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    for ri, (phase, timeline, duration, desc, notes) in enumerate(phases, start=1):
        bg = "E8F4FD" if ri % 2 else "FFFFFF"
        for ci, val in enumerate([phase, timeline, duration, desc, notes]):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
